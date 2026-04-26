# =========================================
# HiddenEdge / SB3PM Advisory & Services Ltd
# =========================================

from flask import Flask, request, jsonify, render_template, send_file, session
from core.cv.cv_service import CVService
import io, os
from docx import Document
import PyPDF2

print("HiddenEdge Engine v1.0 | SB3PM")

# ✅ IMPORTANT: static + templates explicitly defined
app = Flask(
    __name__,
    template_folder="templates",
    static_folder="static"
)

app.secret_key = "hiddenedge_dev_secret"

cv_service = CVService()


# =========================================
# ROUTES
# =========================================

@app.route("/")
def index():
    return render_template("index.html")


# ✅ FIXED: proper Flask rendering (NO file open hacks)
@app.route("/app")
def app_page():
    return render_template("app.html")


@app.route("/eula")
def eula():
    return render_template("eula.html")


@app.route("/email")
def email():
    return render_template("email.html")


@app.route("/submit-email", methods=["POST"])
def submit_email():
    data = request.get_json()
    email = data.get("email")

    if not email:
        return jsonify({"success": False}), 400

    session["user_email"] = email
    return jsonify({"success": True, "redirect": "/app"})


# =========================================
# FILE EXTRACTION
# =========================================

def extract_text_from_docx(file_bytes):
    doc = Document(io.BytesIO(file_bytes))
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])


def extract_text_from_pdf(file_bytes):
    reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
    return "\n".join([p.extract_text() or "" for p in reader.pages])


def extract_text(filename, file_bytes):
    if filename.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    if filename.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    return ""


# =========================================
# ANALYZE
# =========================================

@app.route("/analyze", methods=["POST"])
def analyze():
    files = request.files.getlist("files")
    job_text = request.form.get("job_text", "")

    texts = []
    for f in files:
        t = extract_text(f.filename.lower(), f.read())
        if t:
            texts.append(t)

    result = cv_service.analyze_cv(texts, job_text)
    return jsonify(result)


# =========================================
# SCORE INTELLIGENCE
# =========================================

@app.route("/evaluate_answers", methods=["POST"])
def evaluate_answers():
    data = request.json

    base_score = int(data.get("base_score", 50))
    answers = data.get("answers", "")

    # Length factor
    length_score = min(15, len(answers) // 30)

    # Keyword intelligence (basic but real)
    keywords = ["impact", "result", "delivered", "improved", "managed"]
    relevance_score = sum([1 for k in keywords if k in answers.lower()])

    improvement = min(25, length_score + relevance_score)

    new_score = min(100, base_score + improvement)

    return jsonify({
        "base_score": base_score,
        "improvement": improvement,
        "new_score": new_score
    })


# =========================================
# REFINE CV (DYNAMIC INPUT)
# =========================================

@app.route("/refine-cv", methods=["POST"])
def refine():
    try:
        files = request.files.getlist("files")
        job_text = request.form.get("job_text", "")
        answers = request.form.get("answers", "")

        texts = []
        for f in files:
            t = extract_text(f.filename.lower(), f.read())
            if t:
                texts.append(t)

        result = cv_service.refine_cv_with_answers(texts, job_text, answers)

        return jsonify(result)

    except Exception as e:
        return jsonify({"error": str(e)}), 500


# =========================================
# DOWNLOAD CV
# =========================================

@app.route("/download_cv", methods=["POST"])
def download_cv():
    data = request.json
    cv_text = data.get("cv_text", "")

    doc = Document()
    for line in cv_text.split("\n"):
        doc.add_paragraph(line)

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    return send_file(
        stream,
        as_attachment=True,
        download_name="HiddenEdge_CV.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


# =========================================
# RUN
# =========================================

if __name__ == "__main__":
    app.run(debug=True)