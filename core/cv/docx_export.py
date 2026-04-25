if __name__ == "__main__":
    print("HiddenEdge Engine v1.0 | SB3PM")
# =========================================
# HiddenEdge / SB3PM Advisory & Services Ltd
# Author: Stephan Bals
# © 2026 SB3PM Advisory & Services Ltd
#
# This code is proprietary and confidential.
# Unauthorized use, distribution, or replication is prohibited.
# =========================================

from docx import Document
from docx.shared import Pt


def generate_docx(cv_text, output_path="tailored_cv.docx"):
    doc = Document()

    lines = cv_text.split("\n")

    for line in lines:
        line = line.strip()

        if not line:
            continue

        # Section titles (ALL CAPS)
        if line.isupper():
            doc.add_heading(line, level=1)

        # Bullet points
        elif line.startswith("•"):
            doc.add_paragraph(line[1:].strip(), style="List Bullet")

        # Normal text
        else:
            p = doc.add_paragraph(line)
            p.style.font.size = Pt(11)

    doc.save(output_path)
    return output_path